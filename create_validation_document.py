"""
Create Word Document with ETL Validation Sample
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.etl_validator import ETLValidator
import pandas as pd

def create_validation_document():
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('ETL Mapping Validation Sample', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph(
        'This document demonstrates the ETL mapping validation process using the ETL Parser tool. '
        'It shows how CSV mapping documents are converted into SQL validation queries that can be '
        'executed on your database to ensure data transformation accuracy.'
    )
    
    doc.add_page_break()
    
    # Section 1: Sample Mapping
    doc.add_heading('1. Sample Mapping Document', 1)
    doc.add_paragraph(
        'Below is a sample CSV mapping that defines how data should be transformed from '
        'source to target table:'
    )
    
    # Read the complex mapping CSV
    df = pd.read_csv('examples/complex_mapping.csv')
    
    # Add table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Source Column'
    header_cells[1].text = 'Target Column'
    header_cells[2].text = 'Transformation'
    header_cells[3].text = 'Is Key'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['source_column']) if pd.notna(row['source_column']) else ''
        row_cells[1].text = str(row['target_column']) if pd.notna(row['target_column']) else ''
        row_cells[2].text = str(row['transformation']) if pd.notna(row['transformation']) else ''
        row_cells[3].text = str(row['is_key']) if pd.notna(row['is_key']) else ''
    
    doc.add_page_break()
    
    # Section 2: Mapping Analysis
    doc.add_heading('2. Mapping Analysis', 1)
    
    validator = ETLValidator('examples/complex_mapping.csv')
    validator.load_mappings()
    summary = validator.get_mapping_summary()
    
    doc.add_paragraph(f"Total Mappings: {summary['total_mappings']}")
    doc.add_paragraph(f"Detected Source Tables: {', '.join(summary.get('detected_source_tables', []))}")
    doc.add_paragraph(f"Source Columns ({len(summary['source_columns'])}): {', '.join(summary['source_columns'])}")
    doc.add_paragraph(f"Target Columns ({len(summary['target_columns'])}): {', '.join(summary['target_columns'])}")
    
    doc.add_heading('Key Transformations:', 2)
    
    transformations = [
        ("Date Conversion", "TO_TIMESTAMP(source_table.order_date, 'YYYY-MM-DD')", "Converts string date to timestamp"),
        ("String Cleaning", "UPPER(TRIM(REGEXP_REPLACE(source_table.product_name, '[^a-zA-Z0-9 ]', '')))", "Removes special characters and converts to uppercase"),
        ("Type Casting", "CAST(source_table.unit_price AS DECIMAL(10,2))", "Ensures proper decimal precision"),
        ("Case Statement", "CASE WHEN source_table.order_status = 'N' THEN 'NEW' WHEN...", "Maps status codes to descriptions"),
        ("Concatenation", "CONCAT_WS(', ', source_table.shipping_address, source_table.city, ...)", "Combines address fields"),
        ("Null Handling", "COALESCE(source_table.discount_amount, 0)", "Replaces NULL values with 0"),
        ("Calculated Field", "(source_table.quantity * source_table.unit_price) - COALESCE(...)", "Computes total amount")
    ]
    
    for name, example, description in transformations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{name}: ").bold = True
        p.add_run(f"{description}\n  Example: ")
        run = p.add_run(example)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Section 3: Generated SQL Queries
    doc.add_heading('3. Generated SQL Validation Queries', 1)
    
    source_tables = summary.get('detected_source_tables', [])
    source_table = source_tables[0] if source_tables else 'source_table'
    target_table = 'orders_fact'
    
    queries = validator.generate_validation_queries(
        source_table=source_table,
        target_table=target_table,
        query_type='both'
    )
    
    # Query 1: Source MINUS Target
    doc.add_heading('3.1 Source MINUS Target Query', 2)
    doc.add_paragraph(
        'This query identifies records that exist in the source table but are missing in the '
        'target table after applying transformations. These represent data that failed to load or transform.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(queries['source_minus_target'])  # Full query, no truncation
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    doc.add_page_break()
    
    # Query 2: Target MINUS Source
    doc.add_heading('3.2 Target MINUS Source Query', 2)
    doc.add_paragraph(
        'This query identifies records that exist in the target table but are not present in the '
        'transformed source. These could represent duplicate loads or data from other sources.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(queries['target_minus_source'])  # Full query, no truncation
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    doc.add_page_break()
    
    # Section 4: How to Use
    doc.add_heading('4. How to Execute Validation Queries', 1)
    
    doc.add_heading('4.1 Prerequisites', 2)
    prerequisites = [
        'Access to both source and target databases',
        'Read permissions on source and target tables',
        'SQL client (DBeaver, pgAdmin, SQL Server Management Studio, etc.)',
        'Understanding of your data schema'
    ]
    for item in prerequisites:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.2 Execution Steps', 2)
    steps = [
        ('Connect to Database', 'Open your SQL client and connect to the database containing both source and target tables.'),
        ('Copy Query', 'Copy the generated SQL query from the ETL Parser tool output.'),
        ('Replace Table Names', 'If your actual table names differ from the defaults, update the table names in the WITH clauses.'),
        ('Execute Query', 'Run the query to see validation results.'),
        ('Analyze Results', 'Review the record_count to see how many discrepancies exist.'),
        ('Investigate Discrepancies', 'Use the DETAIL rows (limited to 100) to see actual records with issues.')
    ]
    
    for i, (step, description) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {step}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('4.3 Sample PostgreSQL Execution', 2)
    
    sample_code = '''-- Connect to PostgreSQL
psql -h localhost -U your_user -d your_database

-- Execute validation query
\\i validation_query.sql

-- Or copy-paste directly
WITH source_transformed AS (
  SELECT ...
)
SELECT * FROM source_transformed
EXCEPT
SELECT * FROM target_table;'''
    
    p = doc.add_paragraph()
    run = p.add_run(sample_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('4.4 Sample MySQL Execution', 2)
    
    mysql_code = '''-- Connect to MySQL
mysql -h localhost -u your_user -p your_database

-- Note: MySQL doesn't support EXCEPT, use LEFT JOIN alternative
SELECT s.*
FROM source_transformed s
LEFT JOIN target_table t
  ON s.key_column = t.key_column
WHERE t.key_column IS NULL;'''
    
    p = doc.add_paragraph()
    run = p.add_run(mysql_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Section 5: Expected Results
    doc.add_heading('5. Interpreting Results', 1)
    
    doc.add_heading('5.1 Perfect Match', 2)
    doc.add_paragraph(
        'If both queries return 0 records, your ETL process is working perfectly! '
        'All source records have been correctly transformed and loaded into the target table.'
    )
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Expected Result Table:')
    
    perfect_result = '''validation_type        | record_count
-----------------------+-------------
SOURCE_MINUS_TARGET    |            0
TARGET_MINUS_SOURCE    |            0'''
    
    p = doc.add_paragraph()
    run = p.add_run(perfect_result)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph('\n')
    doc.add_paragraph(
        'When both validation_type queries return 0 record_count, '
        'this indicates perfect data synchronization between source and target.'
    )
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Screenshot: SQL Server Management Studio (SSMS) showing perfect match result:')
    try:
        doc.add_picture('screenshots/01_SSMS_Perfect_Match.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Screenshot not found: {e}]')
        run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph('\n')
    
    doc.add_heading('5.2 Discrepancies Found - Target Has Extra Records', 2)
    doc.add_paragraph(
        'If TARGET_MINUS_SOURCE returns records, it means the target table has records that '
        'are not present in the source after transformation:'
    )
    
    target_extra = '''validation_type        | record_count
-----------------------+-------------
SOURCE_MINUS_TARGET    |            0
TARGET_MINUS_SOURCE    |            5
DETAIL                 |         NULL'''
    
    p = doc.add_paragraph()
    run = p.add_run(target_extra)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(200, 100, 0)
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Screenshot: DBeaver showing target has 5 extra records:')
    try:
        doc.add_picture('screenshots/04_DBeaver_Target_Extra_Records.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Screenshot not found: {e}]')
        run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph('\n')
    
    doc.add_heading('5.3 Discrepancies Found - Source Has Extra Records', 2)
    doc.add_paragraph(
        'If SOURCE_MINUS_TARGET returns records, it means the source has records that '
        'failed to load or transform into the target:'
    )
    
    source_extra = '''validation_type        | record_count
-----------------------+-------------
SOURCE_MINUS_TARGET    |           12
TARGET_MINUS_SOURCE    |            0
DETAIL                 |         NULL'''
    
    p = doc.add_paragraph()
    run = p.add_run(source_extra)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(200, 100, 0)
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Screenshot: SQL Server Management Studio (SSMS) showing source has 12 extra records:')
    try:
        doc.add_picture('screenshots/02_SSMS_Source_Extra_Records.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Screenshot not found: {e}]')
        run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph('\n')
    
    doc.add_heading('5.4 Detailed Discrepancy Records', 2)
    doc.add_paragraph(
        'The queries also return detailed records (up to 100) showing exactly which data '
        'is missing or extra. This helps you investigate the root cause:'
    )
    
    detail_example = '''order_key | customer_key | order_date  | total_amount | status
----------+--------------+-------------+--------------+------------
10001     | C001         | 2024-01-15  |       150.00 | NEW
10005     | C012         | 2024-01-16  |       275.50 | PROCESSING
10008     | C008         | 2024-01-17  |        89.99 | COMPLETED
...       | ...          | ...         |          ... | ...
(12 rows)'''
    
    p = doc.add_paragraph()
    run = p.add_run(detail_example)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_paragraph('\n')
    doc.add_paragraph(
        'Each row shows the actual data that exists in one table but not the other, '
        'allowing you to trace back through your ETL logs to find the issue.'
    )
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Screenshot: pgAdmin showing detailed discrepancy records (12 missing rows):')
    try:
        doc.add_picture('screenshots/05_pgAdmin_Discrepancy_Details.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Screenshot not found: {e}]')
        run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_page_break()
    
    doc.add_heading('5.5 Database Tool Examples', 1)
    
    doc.add_heading('SQL Server Management Studio (SSMS)', 2)
    doc.add_paragraph('To execute in SSMS:')
    ssms_steps = [
        'Open SQL Server Management Studio and connect to your database server',
        'Click "New Query" to open a query window',
        'Paste the generated validation query',
        'Press F5 or click "Execute" to run the query',
        'View results in the "Results" pane at the bottom',
        'If record_count = 0, you\'ll see a green checkmark indicating success'
    ]
    for step in ssms_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    run = p.add_run('📸 INSERT SCREENSHOT HERE: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = p.add_run('Paste your actual SSMS screenshot showing the query execution and results pane.')
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph('\n')
    
    doc.add_heading('DBeaver', 2)
    doc.add_paragraph('To execute in DBeaver:')
    dbeaver_steps = [
        'Open DBeaver and connect to your database',
        'Right-click on your database connection and select "SQL Editor" → "New SQL Script"',
        'Paste the generated validation query',
        'Press Ctrl+Enter (or click the "Execute SQL Statement" button)',
        'View results in the "Data" tab below the editor',
        'Use the export button to save results to CSV/Excel if needed'
    ]
    for step in dbeaver_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Screenshot: DBeaver showing perfect validation match:')
    try:
        doc.add_picture('screenshots/03_DBeaver_Perfect_Match.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[Screenshot not found: {e}]')
        run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph('\n')
    
    doc.add_heading('pgAdmin (PostgreSQL)', 2)
    doc.add_paragraph('To execute in pgAdmin:')
    pgadmin_steps = [
        'Open pgAdmin and connect to your PostgreSQL server',
        'Navigate to your database in the tree view',
        'Click "Tools" → "Query Tool"',
        'Paste the generated validation query',
        'Click the "Execute" button (▶) or press F5',
        'View results in the "Data Output" tab',
        'Check the message tab for execution summary'
    ]
    for step in pgadmin_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Note: See pgAdmin screenshot above (Section 5.4) showing detailed discrepancy records.')
    
    doc.add_page_break()
    
    doc.add_heading('5.6 Common Discrepancy Scenarios', 1)
    
    issues = [
        ('Source > Target', 'Source has more records than target', [
            'ETL job failed midway',
            'Filter conditions in ETL removing valid records',
            'Transformation errors causing records to be rejected',
            'Target table constraints (primary key violations)'
        ]),
        ('Target > Source', 'Target has more records than source', [
            'Multiple ETL runs loading same data (duplicates)',
            'Target table receiving data from multiple sources',
            'Historical data in target not in current source',
            'Manual data inserts in target'
        ])
    ]
    
    for title, description, causes in issues:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
        
        doc.add_paragraph('Common causes:', style='List Bullet')
        for cause in causes:
            doc.add_paragraph(cause, style='List Bullet 2')
    
    doc.add_page_break()
    
    # Section 6: Best Practices
    doc.add_heading('6. Best Practices', 1)
    
    practices = [
        ('Run Before and After ETL', 'Execute validation queries before ETL starts and after it completes to confirm successful processing.'),
        ('Schedule Regular Validations', 'Set up automated jobs to run these queries daily/weekly to catch data quality issues early.'),
        ('Keep Query History', 'Save query results over time to track data quality trends and identify recurring issues.'),
        ('Document Exceptions', 'Some discrepancies may be expected (e.g., filtered records). Document these business rules.'),
        ('Alert on Thresholds', 'Set up alerts when discrepancy count exceeds acceptable thresholds.'),
        ('Test with Subsets First', 'When working with large tables, test queries on date-partitioned subsets first.'),
        ('Version Control Mappings', 'Keep your CSV mapping documents in version control (Git) to track changes over time.'),
        ('Peer Review Changes', 'Have another team member review mapping changes before deployment.')
    ]
    
    for title, description in practices:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # Section 7: Tool Features
    doc.add_heading('7. ETL Parser Tool Features', 1)
    
    features = [
        ('Multiple File Formats', 'Supports CSV and Excel (.xlsx, .xls) mapping documents'),
        ('Auto Source Detection', 'Automatically detects source table names from transformation column'),
        ('Multiple Database Types', 'Generate queries for PostgreSQL, MySQL, Oracle, SQL Server, Snowflake'),
        ('AI Enhancement', 'Optional AI-powered query optimization and transformation suggestions'),
        ('Bidirectional Validation', 'Generates both Source→Target and Target→Source comparison queries'),
        ('Web Interface', 'User-friendly web interface - no coding required'),
        ('Free to Use', 'Open source and completely free for commercial and personal use')
    ]
    
    for feature, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph('\n')
    doc.add_heading('Contact & Support', 2)
    doc.add_paragraph('Tool URL: https://etl-parser.onrender.com')
    doc.add_paragraph('GitHub: https://github.com/hkrishnan62/ETL_Parser')
    doc.add_paragraph('For issues or feature requests, please open a GitHub issue.')
    
    # Save document
    filename = 'ETL_Validation_Sample.docx'
    doc.save(filename)
    print(f"✅ Document created successfully: {filename}")
    print(f"📄 Document contains {len(doc.paragraphs)} paragraphs across multiple sections")
    print(f"📊 Includes sample mapping with {len(df)} transformations")
    print(f"🔍 Contains generated SQL validation queries")

if __name__ == '__main__':
    create_validation_document()
